from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path

from PIL import Image
from pytesseract import TesseractNotFoundError, image_to_string


@dataclass(frozen=True)
class ExtractedText:
    text: str
    metadata: dict[str, int]
    warnings: tuple[str, ...] = ()


class OCRNotAvailableError(RuntimeError):
    """Raised when OCR is requested but Tesseract is unavailable or disabled."""


def _normalize_text(text: str, max_chars: int | None) -> tuple[str, list[str]]:
    """Collapse whitespace and optionally truncate. Returns (text, warnings)."""
    warnings: list[str] = []
    text = re.sub(r"\s+", " ", text).strip()
    if max_chars is not None and len(text) > max_chars:
        text = text[:max_chars].rsplit(" ", 1)[0].strip() or text[:max_chars]
        warnings.append(f"text_truncated_to_{max_chars}_chars")
    return text, warnings


class FileTextExtractor:
    def __init__(
        self,
        *,
        ocr_enabled: bool = True,
        tesseract_lang: str = "rus+eng",
        max_chars: int | None = 200_000,
        max_pages: int | None = 50,
    ) -> None:
        self._ocr_enabled = ocr_enabled
        self._tesseract_lang = tesseract_lang or "eng"
        self._max_chars = max_chars
        self._max_pages = max_pages

    def extract(
        self,
        *,
        path: Path,
        file_type: str,
        max_chars: int | None = None,
        max_pages: int | None = None,
    ) -> ExtractedText:
        limit_chars = max_chars if max_chars is not None else self._max_chars
        limit_pages = max_pages if max_pages is not None else self._max_pages
        if file_type == "pdf":
            return self._extract_pdf(path, limit_chars=limit_chars, limit_pages=limit_pages)
        if file_type == "docx":
            return self._extract_docx(path, limit_chars=limit_chars)
        if file_type == "image":
            return self._extract_image(path, limit_chars=limit_chars)
        raise ValueError(f"Unsupported file type: {file_type}")

    def extract_from_bytes(
        self,
        data: bytes,
        file_type: str,
        filename: str = "",
        max_chars: int | None = None,
        max_pages: int | None = None,
    ) -> ExtractedText:
        limit_chars = max_chars if max_chars is not None else self._max_chars
        limit_pages = max_pages if max_pages is not None else self._max_pages
        if file_type == "pdf":
            return self._extract_pdf_bytes(data, limit_chars=limit_chars, limit_pages=limit_pages)
        if file_type == "docx":
            return self._extract_docx_bytes(data, limit_chars=limit_chars)
        if file_type == "image":
            return self._extract_image_bytes(data, limit_chars=limit_chars)
        raise ValueError(f"Unsupported file type: {file_type}")

    def _extract_pdf(
        self,
        path: Path,
        *,
        limit_chars: int | None,
        limit_pages: int | None,
    ) -> ExtractedText:
        try:
            from pypdf import PdfReader
        except ModuleNotFoundError as exc:
            raise ImportError("pypdf is required to extract text from PDF files.") from exc
        reader = PdfReader(str(path))
        total_pages = len(reader.pages)
        pages_to_use = total_pages
        if limit_pages is not None and total_pages > limit_pages:
            pages_to_use = limit_pages
        pages_text: list[str] = []
        for i in range(pages_to_use):
            try:
                page = reader.pages[i]
                pages_text.append(page.extract_text() or "")
            except Exception:
                pages_text.append("")
        text = "\n".join(pages_text).strip()
        all_warnings: list[str] = []
        if limit_pages is not None and total_pages > limit_pages:
            all_warnings.append(f"pages_limited_{total_pages}_to_{limit_pages}")
        text, norm_warnings = _normalize_text(text, limit_chars)
        all_warnings.extend(norm_warnings)
        return ExtractedText(
            text=text,
            metadata={"pages": pages_to_use, "characters": len(text)},
            warnings=tuple(all_warnings),
        )

    def _extract_pdf_bytes(
        self,
        data: bytes,
        *,
        limit_chars: int | None,
        limit_pages: int | None,
    ) -> ExtractedText:
        try:
            from pypdf import PdfReader
        except ModuleNotFoundError as exc:
            raise ImportError("pypdf is required to extract text from PDF files.") from exc
        reader = PdfReader(io.BytesIO(data))
        total_pages = len(reader.pages)
        pages_to_use = total_pages
        if limit_pages is not None and total_pages > limit_pages:
            pages_to_use = limit_pages
        pages_text = []
        for i in range(pages_to_use):
            try:
                page = reader.pages[i]
                pages_text.append(page.extract_text() or "")
            except Exception:
                pages_text.append("")
        text = "\n".join(pages_text).strip()
        all_warnings: list[str] = []
        if limit_pages is not None and total_pages > limit_pages:
            all_warnings.append(f"pages_limited_{total_pages}_to_{limit_pages}")
        text, norm_warnings = _normalize_text(text, limit_chars)
        all_warnings.extend(norm_warnings)
        return ExtractedText(
            text=text,
            metadata={"pages": pages_to_use, "characters": len(text)},
            warnings=tuple(all_warnings),
        )

    def _extract_docx(self, path: Path, *, limit_chars: int | None) -> ExtractedText:
        try:
            from docx import Document
        except ModuleNotFoundError as exc:
            raise ImportError("python-docx is required to extract text from DOCX files.") from exc
        document = Document(str(path))
        parts: list[str] = []
        for p in document.paragraphs:
            if p.text:
                parts.append(p.text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        text = "\n".join(parts).strip()
        text, warnings = _normalize_text(text, limit_chars)
        return ExtractedText(
            text=text,
            metadata={"paragraphs": len(document.paragraphs), "characters": len(text)},
            warnings=tuple(warnings),
        )

    def _extract_docx_bytes(self, data: bytes, *, limit_chars: int | None) -> ExtractedText:
        try:
            from docx import Document
        except ModuleNotFoundError as exc:
            raise ImportError("python-docx is required to extract text from DOCX files.") from exc
        document = Document(io.BytesIO(data))
        parts = []
        for p in document.paragraphs:
            if p.text:
                parts.append(p.text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        text = "\n".join(parts).strip()
        text, warnings = _normalize_text(text, limit_chars)
        return ExtractedText(
            text=text,
            metadata={"paragraphs": len(document.paragraphs), "characters": len(text)},
            warnings=tuple(warnings),
        )

    def _extract_image(self, path: Path, *, limit_chars: int | None) -> ExtractedText:
        if not self._ocr_enabled:
            raise OCRNotAvailableError("OCR disabled via configuration.")
        try:
            image = Image.open(path)
            image.load()
            text = image_to_string(image, lang=self._tesseract_lang) or ""
        except TesseractNotFoundError as exc:
            raise OCRNotAvailableError("Tesseract OCR is not available.") from exc
        text = text.strip()
        text, warnings = _normalize_text(text, limit_chars)
        return ExtractedText(
            text=text,
            metadata={"characters": len(text)},
            warnings=tuple(warnings),
        )

    def _extract_image_bytes(self, data: bytes, *, limit_chars: int | None) -> ExtractedText:
        if not self._ocr_enabled:
            raise OCRNotAvailableError("OCR disabled via configuration.")
        try:
            image = Image.open(io.BytesIO(data))
            image.load()
            text = image_to_string(image, lang=self._tesseract_lang) or ""
        except TesseractNotFoundError as exc:
            raise OCRNotAvailableError("Tesseract OCR is not available.") from exc
        text = text.strip()
        text, warnings = _normalize_text(text, limit_chars)
        return ExtractedText(
            text=text,
            metadata={"characters": len(text)},
            warnings=tuple(warnings),
        )
